import os
import json
import pandas as pd
from pathlib import Path
from typing import List, Dict, Any, Optional
from config import PROCESSED_DIR, DATA_DIR

class DocumentImporter:
    """Multi-format Document Importer with Smart Incremental Caching & Deduplication."""

    CACHE_FILE = DATA_DIR / "gdrive_parse_cache.json"

    def __init__(self):
        self.gdrive_root = self._detect_gdrive_root()
        self.cache: Dict[str, float] = self._load_cache()

    def _load_cache(self) -> Dict[str, float]:
        if self.CACHE_FILE.exists():
            try:
                with open(self.CACHE_FILE, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                pass
        return {}

    def _save_cache(self):
        try:
            with open(self.CACHE_FILE, "w", encoding="utf-8") as f:
                json.dump(self.cache, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def _detect_gdrive_root(self) -> Path:
        cloud_storage = Path.home() / "Library" / "CloudStorage"
        if cloud_storage.exists():
            gdirs = list(cloud_storage.glob("GoogleDrive-*"))
            if gdirs:
                my_drive = gdirs[0] / "我的雲端硬碟"
                if not my_drive.exists():
                    my_drive = gdirs[0] / "My Drive"
                if not my_drive.exists():
                    my_drive = gdirs[0]
                return my_drive
        return Path.home()

    def extract_text_from_pdf(self, pdf_path: Path) -> str:
        text_parts = []
        try:
            import pdfplumber
            with pdfplumber.open(str(pdf_path)) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    tables = page.extract_tables()
                    table_str = ""
                    if tables:
                        for table in tables:
                            try:
                                df_tbl = pd.DataFrame(table[1:], columns=table[0])
                                table_str += "\n\n" + df_tbl.to_markdown(index=False) + "\n\n"
                            except Exception:
                                pass
                    combined = (page_text + "\n" + table_str).strip()
                    if combined:
                        text_parts.append(f"--- Page {i+1} ---\n{combined}")
        except Exception as e:
            print(f"[pdfplumber] Warning: {e}, falling back to pypdf...")

        if not text_parts:
            try:
                from pypdf import PdfReader
                reader = PdfReader(str(pdf_path))
                for i, page in enumerate(reader.pages):
                    txt = page.extract_text()
                    if txt and txt.strip():
                        text_parts.append(f"--- Page {i+1} ---\n{txt}")
            except Exception as e:
                text_parts.append(f"[Error extracting PDF text: {e}]")

        return "\n\n".join(text_parts)

    def extract_text_from_xlsx(self, xlsx_path: Path) -> str:
        text_parts = []
        try:
            excel_file = pd.ExcelFile(str(xlsx_path))
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                if not df.empty:
                    table_md = df.to_markdown(index=False)
                    text_parts.append(f"### Sheet WorkBook: `{sheet_name}`\n\n{table_md}")
        except Exception as e:
            text_parts.append(f"[Error parsing Excel file: {e}]")

        return "\n\n".join(text_parts)

    def extract_text_from_csv(self, csv_path: Path) -> str:
        try:
            df = pd.read_csv(str(csv_path))
            return df.to_markdown(index=False)
        except Exception as e:
            return f"[Error parsing CSV: {e}]"

    def extract_text_from_docx(self, docx_path: Path) -> str:
        try:
            import docx
            doc = docx.Document(str(docx_path))
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())
            for t in doc.tables:
                table_data = []
                for row in t.rows:
                    table_data.append([cell.text.strip() for cell in row.cells])
                if table_data and len(table_data) > 1:
                    df = pd.DataFrame(table_data[1:], columns=table_data[0])
                    parts.append(df.to_markdown(index=False))
            return "\n\n".join(parts)
        except Exception as e:
            return f"[Error parsing Word doc: {e}]"

    def import_all_gdrive_documents(self) -> List[Path]:
        """Scan Google Drive folders for supported formats with deduplication & caching."""
        imported_files = []

        search_dirs = [
            Path.home() / "Documents" / "Google Drive",
            self.gdrive_root / "ChronoAFR_Sync",
            self.gdrive_root
        ]
        cache_updated = False
        processed_filenames = set()

        for s_dir in search_dirs:
            if not s_dir.exists():
                continue
            for file in s_dir.iterdir():
                if file.is_dir() or file.name.startswith("."):
                    continue

                fname = file.name
                # Skip system-generated reports or converted GDrive_ files to prevent loops!
                if fname.startswith("GDrive_") or fname.endswith("_Report.md") or fname.endswith("_Report.txt"):
                    continue

                # Deduplicate if file with same name was already processed in this scan cycle
                if fname in processed_filenames:
                    continue

                ext = file.suffix.lower()
                # Include full original filename in processed name: GDrive_Amazon-2025-Annual-Report.pdf.md
                out_path = PROCESSED_DIR / f"GDrive_{file.name}.md"
                mtime = file.stat().st_mtime

                processed_filenames.add(fname)

                # Incremental Check: If out_path exists and file mtime is unchanged in cache, skip re-parsing!
                if out_path.exists() and self.cache.get(fname) == mtime:
                    imported_files.append(out_path)
                    continue

                text = ""
                if ext == ".pdf":
                    text = self.extract_text_from_pdf(file)
                elif ext in [".xlsx", ".xls"]:
                    text = self.extract_text_from_xlsx(file)
                elif ext == ".csv":
                    text = self.extract_text_from_csv(file)
                elif ext in [".docx", ".doc"]:
                    text = self.extract_text_from_docx(file)
                elif ext in [".txt", ".md"]:
                    try:
                        text = file.read_text(encoding="utf-8", errors="ignore")
                    except Exception:
                        text = ""

                if text and len(text.strip()) > 10:
                    md_content = f"# [Google Drive Document] {file.name}\n\n"
                    md_content += f"- **檔案名稱**: `{file.name}`\n"
                    md_content += f"- **格式類別**: `{ext.upper()}`\n"
                    md_content += f"- **檔案大小**: `{file.stat().st_size / 1024:.1f} KB`\n"
                    md_content += f"- **同步狀態**: 已成功轉譯寫入 ChronoAFR 資料庫\n\n"
                    md_content += "## 內容內文與數據表格\n\n" + text

                    with open(out_path, "w", encoding="utf-8") as f:
                        f.write(md_content)

                    self.cache[fname] = mtime
                    cache_updated = True
                    imported_files.append(out_path)

        if cache_updated:
            self._save_cache()

        return imported_files
